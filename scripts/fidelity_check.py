import json, glob, re
from docx import Document
from docx.oxml.ns import qn

PATH = r"C:\Users\Administrator\Downloads\换乘恋爱_全文整理版.docx"
DATA = r"C:\Users\Administrator\WorkBuddy\2026-08-12-14-38-27\app\src\data"

doc = Document(PATH)

# 1. Collect bookmarked content paragraphs in document order
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
META_STYLES = {'Title','Subtitle','Heading 1','Module Label','Production Note'}
content_bookmarks = {}  # bm name -> paragraph index
para_bm = {}  # para idx -> bm name
for pidx, p in enumerate(doc.paragraphs):
    bm_elems = p._p.findall('.//' + qn('w:bookmarkStart'))
    for b in bm_elems:
        name = b.get(qn('w:name'))
        if name and re.match(r'P\d{2}_\d{4}', name):
            content_bookmarks[name] = pidx
            para_bm[pidx] = name

# ordered list of (bm, text)
doc_order = []
seen = set()
for pidx, p in enumerate(doc.paragraphs):
    if pidx in para_bm:
        bm = para_bm[pidx]
        if bm in seen: 
            continue
        seen.add(bm)
        doc_order.append((bm, p.text))

print("DOCX content blocks (bookmarked):", len(doc_order))

# 2. Load parsed JSON
all_blocks = []
for f in sorted(glob.glob(DATA + "/chapter-*.json")):
    for b in json.load(open(f, encoding='utf-8')):
        all_blocks.append(b)
print("Parsed blocks:", len(all_blocks))

# 3. Compare in order
mismatch = 0
missing_source = 0
text_diff = 0
missing_in_json = 0
# index by sourceId
json_by_sid = {b['sourceId']: b for b in all_blocks}

i = 0  # doc pointer
j = 0  # json pointer
order_ok = True
detail = []
for (bm, txt) in doc_order:
    if i >= len(all_blocks):
        missing_in_json += 1
        detail.append(("MISSING_IN_JSON", bm, txt[:40]))
        continue
    cur = all_blocks[i]
    if cur['sourceId'] != bm:
        order_ok = False
        detail.append(("ORDER", cur['sourceId'], bm))
        # try to find it later
        if bm in json_by_sid:
            cur = json_by_sid[bm]
        else:
            missing_in_json += 1
            i += 1
            continue
    if cur['text'] != txt:
        text_diff += 1
        if text_diff <= 15:
            detail.append(("TEXTDIFF", bm, repr(cur['text'][:60]), repr(txt[:60])))
    i += 1

# also check duplicate sourceIds
from collections import Counter
sidc = Counter(b['sourceId'] for b in all_blocks)
dup = [s for s,c in sidc.items() if c>1]

print("\n=== FIDELITY REPORT ===")
print("DOCX bookmarked blocks :", len(doc_order))
print("Parsed blocks          :", len(all_blocks))
print("Order consistent       :", order_ok)
print("Text differences       :", text_diff)
print("Duplicate sourceIds    :", len(dup))
print("Missing in JSON (more in doc than parsed):", missing_in_json)
print("\n--- sample diffs (first 15) ---")
for d in detail[:15]:
    print(d)
